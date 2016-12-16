import cherrypy
import random
from cherrypy.lib.static import serve_file
import os
from io import StringIO
from subprocess import call
import sys
import dill as pickle
#import pickle

try:
    from sympy import latex
except:
    print('No sympy installed')

def section(text):
    print("<h1>"+text+"</h1>")

#Globals = {'section':section}
#Locals = {}

def genRandomStr(length):
    chars = 'qwertyuiopasdfghjklzxcvbnm1234567890QWERTYUIOPASDFGHJKLZXCVBNM'
    randomStr = ''.join(random.choice(chars) for i in range(length))
    return randomStr

def getAllInside(first,last,content):
    # returns a dict of keys first+value+last and the values
    valuesDict = {}
    for i in range(content.count(first)):
        try:
            f = content.index(first)
            l = content[f+len(first):].index(last)
            key = content[f:f+len(first)+l+len(last)]
            value = key.replace(first,'').replace(last,'')
        except ValueError:
            print('Error')
            continue
        valuesDict[key] = value
        content = content.replace(key,'')
    return valuesDict

def getInside(first,last,content):
    f = content.index(first)
    l = content[f+len(first):].index(last)
    key = content[f:f+len(first)+l+len(last)]
    value = key.replace(first,'').replace(last,'')
    return value

def startWith(text,content):
    return content.replace('\n','').replace(' ','')[:len(text)] == text

def emptyLine(content):
    if content.replace('\n','').replace(' ','') == '':
        return True
    else:
        return False

class Capturing(list):
    def __enter__(self):
        self._stdout = sys.stdout
        sys.stdout = self._stringio = StringIO()
        return self
    def __exit__(self, *args):
        self.extend(self._stringio.getvalue().splitlines())
        print(self)
        sys.stdout = self._stdout
	  
class WillNotebook(object):
    emptyLineSymbol = '.'
    archive = {}
    references = {}
    @cherrypy.expose
    def index(self):
        docID= genRandomStr(8)
        while docID in self.archive:
            docID = genRandomStr(8)
        notebook = open('notebook.html','r')
        notebook = notebook.read().replace('!@docID@!',docID)
        #The globals can be used to perform Python default functions for WillNotebook
        #It is ignored in the save state, so only locals is saved in the document
        self.archive[docID] = {'Globals':{'section':section},'Locals':{},'page':[]}
        self.references[docID] = {'keys':{},'References':''}
        return notebook

    @cherrypy.expose
    def newCell(self,docID,index):
        print('New cell called')
        self.archive[docID]['page'].insert(int(index),{'content':'','output':'.'})
        return 'Cell inserted'

    @cherrypy.expose
    def deleteCell(self,docID,index):
        index = int(index)
        print('Delete cell called')
        print('Index: ',index)
        if 'type' in self.archive[docID]['page'][index]['content']:
            if self.archive[docID]['page'][index]['content']['type'] == 'image':
                filename = self.archive[docID]['page'][index]['content']['img']
                os.remove(os.getcwd()+'/Archieves/Images/'+filename)
        self.archive[docID]['page'].pop(int(index))
        return 'Cell deleted'

    @cherrypy.expose
    def evalCell(self,docID,cell,content):
        cell = int(cell)
        print(cell)
        if cell == len(self.archive[docID]['page']):
            self.archive[docID]['page'].append({'content':content,'output':'.'})
        self.archive[docID]['page'][cell]['content'] = content
        if startWith('#code',content):
            output = self.handlePythonCode(docID,content)
            ## {{}} soh nao sera avaliado em #code
        else:
            if '{{' in content and '}}' in content:
                content = self.handleValues(docID,content)
            ## Order is important here, first the ***
            if '***' in content:
                print('BoldItalic parts')
                content = self.handleBoldItalics(content)
            if '**' in content:
                print('Bold parts')
                content = self.handleBold(content)
            if '*' in content:
                print('Italic parts')
                content = self.handleItalics(content)
            if '\cite{' in content and '}' in content:
                content = self.handleCitations(docID,content)
            if startWith('!#',content):
                output = self.handleSections(content)
            elif startWith('!eq',content):
                output = self.handleEquations(content)
            elif startWith('!title',content):
                output = self.handleTitle(content)
            elif startWith('!ref',content):
                output = self.handleReferences(docID,content)
            else:
                output = content
        if emptyLine(output):
            output = self.emptyLineSymbol
        self.archive[docID]['page'][cell]['output'] = output
        print('Out: ',output)
        return output

    def handlePythonCode(self,docID,content):
        with Capturing() as output:
            try:
                exec(content,self.archive[docID]['Globals'],self.archive[docID]['Locals'])
            except Exception as e:
                print('Exception: ', e)
        if output == []:
            label = ''
            for line in content.split('\n'):
                if '#code' in line:
                    label = line.replace('#code','')
                    break
            return '<font class="dontprint" color="green">[code]'+label+'</font>'
        else:
            # Retorna o primeiro indice da lista (str)
            return output[0]

    def handleValues(self,docID,content):
        toEvaluate = getAllInside('{{','}}',content)
        for expr in toEvaluate:
            try:
                out = eval(toEvaluate[expr],self.archive[docID]['Globals'],self.archive[docID]['Locals'])
                objType = str(type(out))
                if 'sympy' in objType or 'list' in objType and 'sympy' in str(type(out[0])):
                    print('Is a sympy object!')
                    try:
                        out = str(latex(out))
                        if not '!eq' in content:
                            out = '$'+out+'$'
                    except:
                        print('Latex error')
                        out = str(out)
                else: 
                    print('Not a sympy object! It it: ',objType)
                    out = str(out)
                content = content.replace(expr,out)
            except Exception as e:
                print('excep: ',e)
                content = 'Exception: '+str(e)
        return content

    def handleItalics(self,content):
        toItalicize = getAllInside('*','*', content)
        for expr in toItalicize:
            print('Italics :',expr)
            italic = '<i>'+toItalicize[expr]+'</i>'
            content = content.replace(expr,italic)
        return content

    def handleBold(self,content):
        toBold = getAllInside('**','**', content)
        for expr in toBold:
            print('Bold :',expr)
            bold = '<b>'+toBold[expr]+'</b>'
            content = content.replace(expr,bold)
        return content

    def handleBoldItalics(self,content):
        toBoldItalicize = getAllInside('***','***', content)
        for expr in toBoldItalicize:
            print('BoldItalics :',expr)
            boldItalic = '<i><b>'+toBoldItalicize[expr]+'</b></i>'
            content = content.replace(expr,boldItalic)
        return content

    def handleCitations(self,docID,content):
        print('Citations working')
        def getRef(line):
            print('Getting from: ',line)
            keyword = 'indent" >'
            i = line.index(keyword)
            reference = line[i+len(keyword):]
            print('The ref is: ',reference)
            while ' ' in reference[0]:
                reference = reference[1:]
            return reference

        def makeReferences(docID):
            '''Updates the references'''
            self.references[docID]['References'] = ''
            ### Make the tex file and compile it ###
            citations = open(os.getcwd()+'/Archieves/cit.tex','w')
            citations.write('''\\documentclass{article}
\\usepackage[T1]{fontenc}
\\usepackage[utf8]{inputenc}
\\usepackage[alf,abnt-etal-list=0,abnt-repeated-author-omit=no]{abntex2cite}
\\usepackage{url}

\\begin{document}
''')
            refs = []
            for ref in self.references[docID]['keys']:
                citations.write(ref+'\n\n')
                refs.append(ref)
            citations.write('''\\bibliography{database}
\end{document}''')
            citations.close()
            call(['htlatex','cit.tex'],cwd=os.getcwd()+'/Archieves/')
            call(['bibtex','cit.aux'],cwd=os.getcwd()+'/Archieves/')
            call(['htlatex','cit.tex'],cwd=os.getcwd()+'/Archieves/')
            ### END: Make the tex file and compile it ###
            ### Parse the compiled HTML file ###
            html = open(os.getcwd()+'/Archieves/cit.html','r',errors='replace',encoding='latin1')
            n = 0
            beginReference = False
            for line in html:
                ### Get citation ###
                if n < len(refs):
                    if 'class="noindent"' in line or 'class="indent"' in line:
                        self.references[docID]['keys'][refs[n]] = getRef(line)
                        n+=1
                ### Get reference ###
                else:
                    ### Order of the if's is important
                    if '</div>' in line:
                        beginReference = False
                    if beginReference:
                        self.references[docID]['References']+= line
                    if 'thebibliography' in line:
                        beginReference = True

            ### END: Parse the compiled HTML file ###

        toCitate = getAllInside('\cite{','}',content)
        ## Check if there is a new citation
        changed = False
        for citation in toCitate:
            if not citation in self.references[docID]['keys']:
                self.references[docID]['keys'][citation] = ''
                changed = True
        if changed:
            makeReferences(docID)
        for citation in toCitate:
            content = content.replace(citation,self.references[docID]['keys'][citation])    
        return content

    def handleReferences(self,docID,content):
        output = '<h1>References</h1>\n'
        output += self.references[docID]['References']
        return output

    def handleSections(self,content):
        n = 1 #number of #'s in the section mark
        pos = content.index('!#')+2
        for char in content[pos:]:
            if char == '#':
                n+=1
            else:
                break
        heading = ''
        for line in content.split('\n'):
            if '!#' in line:
                label = line.replace('!'+'#'*n,'')
                if label:
                    if label[0] == ' ':
                        label = label[1:]
            else:
                heading += line
        if label:
            return '<h'+str(n)+' id="'+label+'">'+heading+'</h'+str(n)+'>'
        else:
            return '<h'+str(n)+'>'+heading+'</h'+str(n)+'>'

    def handleEquations(self,content):
        eqContent = ''
        for line in content.split('\n'):
            if '!eq' in line:
                label = line.replace('!eq ','')
            else:
                eqContent += line
        eq = '\\begin{equation}'
        if label.replace('<label>',''):
            eq += '\label{'+label+'}\n'
        eq += eqContent+'\end{equation}'
        return eq

    def handleTitle(self,content):
        title = content.replace('!title ','')
        return '<center><b><font size="7">'+title+'</font></b></center><br><br>'
    
    @cherrypy.expose
    def image(self,docID,cell,img,label,source,caption):
        cell = int(cell)
        filename = img.filename
        if cell == len(self.archive[docID]['page']):
            self.archive[docID]['page'].append({'content':{'type':'image','img':filename,'label':label,'source':source,'caption':caption},'output':'.'})
        i = open(os.getcwd()+'/Archieves/Images/'+filename,'wb')
        while True:
            data = img.file.read(4096)
            if not data:
                break
            else:
                i.write(data)
            print('Loading...')
        i.close()
        if label:
            output = '<br><center><figcaption id="'+label+'">'+caption+'</figcaption>'+'<img style="max-width:800px" src="Archieves/Images/'+filename+'"><br>Source: '+source+'</center><br>'
        else:
            output = '<br><center><figcaption>'+caption+'</figcaption>'+'<img style="max-width:800px" src="Archieves/Images/'+filename+'"><br>Source: '+source+'</center><br>'
        self.archive[docID]['page'][cell]['output'] = output
        return output

    @cherrypy.expose
    def saveFile(self,docID,filename,extension):
        if extension == 'will':
            self.saveAsWill(docID,filename)
        elif extension == 'tex':
            self.saveAs(docID,filename,extension)
        elif extension == 'docx':
            self.saveAs(docID,filename,extension)
        elif extension == 'pdflatex':
            self.saveAs(docID,filename,extension)
            extension = 'pdf'
        return serve_file(os.getcwd()+'/Archieves/'+filename+'.'+extension,"application/x-download","attachment")

    def saveAsWill(self,docID,filename):
        archive = open(os.getcwd()+'/Archieves/'+filename+'.will','wb')
        try:
            pickle.dump(self.archive[docID],archive)
        except Exception as e:
            print('Could not save state! ',e)
            self.archive[docID]['Locals'] = {}
            pickle.dump(self.archive[docID],archive)
        archive.close()

    def saveAs(self,docID,filename,docFormat,article=True):
        if article:
            texClass = 'article'
        else:
            texClass = 'report'
        if docFormat == 'tex':
            from texExporter import TexExporter
            exporter = TexExporter(filename,texClass)
        elif docFormat == 'pdflatex':
            from pdfLatexExporter import PdfLatexExporter
            exporter = PdfLatexExporter(filename,texClass)
        elif docFormat == 'docx':
            from docxExporter import DocxExporter
            exporter = DocxExporter(filename,texClass)

        for cell in self.archive[docID]['page']:
            content = cell['content']
            show = True
            if '<h' in cell['output']:
                for level in range(1,6):
                    n = str(level)
                    if '<h'+n in cell['output']:
                        label = ''
                        endTag = '>'
                        if 'id="' in cell['output']:
                            label = getInside('id="','"',cell['output'])
                            endTag = ' id="'+label+'">'
                        title = next(iter(getAllInside('<h'+n+endTag,'</h'+n+'>',cell['output']).values()))
                        exporter.addHeading(title,level,label)
                        break
            elif 'class="dontprint"' in cell['output']:
                show = False
            ### special cells ###
            elif 'type' in content and not type(content) == str:
                if content['type'] == 'image':
                    img = content['img']
                    caption = content['caption']
                    source = content['source']
                    label = content['label']
                    exporter.addFigure(img,caption,source,label)
                else:
                    raise NotImplemented
            else:
                if show:
                    exporter.addText(cell['output'])

        print('ta no close')
        exporter.close()
        print('Acabou')

    def saveAsDocx(self,docID,filename):
        from docx import Document
        from docx.shared import Inches
        d = Document()
        for cell in self.archive[docID]['page']:
            content = cell['content']
            if '!# ' in content:
                d.add_heading(content.replace('!# ',''),level=1)
            elif '!## ' in content:
                d.add_heading(content.replace('!## ',''),level=2)
            elif '!### ' in content:
                d.add_heading(content.replace('!### ',''),level=3)
            elif '!#### ' in content:
                d.add_heading(content.replace('!#### ',''),level=4)
            elif '!##### ' in content:
                d.add_heading(content.replace('!##### ',''),level=5)
            else:
                d.add_paragraph(content)

        d.save(os.getcwd()+'/Archieves/'+filename+'.docx')

    @cherrypy.expose
    def open(self,docID,toOpen):
        filename = toOpen.filename
        d = open(os.getcwd()+'/Archieves/'+filename,'wb')
        while True:
            data = toOpen.file.read(4096)
            if not data:
                break
            else:
                d.write(data)
            print('Loading...')
        d.close()
        req = self.openFile(docID,filename)
        return req

    @cherrypy.expose
    def openFile(self,docID,filename):
        self.archive[docID] = {}
        print('Openning file ',filename)
        try:
            archive = open(os.getcwd()+'/Archieves/'+filename,'rb')
        except:
            content = 'Error loading file. File not found.'
            output = content
            print('Error opening file. File does not exist!')
            return '<center id="c1"><textarea id="1" action="evalCell" style="width: 800px; display: none;">'+content+'</textarea></center><center id="co1"><div id="o1" class="paragraph">'+output+'</div></center>'
        self.archive[docID] = pickle.load(archive)
        self.archive[docID]['Globals'] = {'section':section}
        archive.close()
        notebook = ''
        for cell,stuff in enumerate(self.archive[docID]['page']):
            # toDo check if cell is 'image'. If so, construct the image, not
            # the textArea
            print(cell)
            if 'type' in stuff['content'] and not type(stuff['content']) == str:
                if stuff['content']['type'] == 'image':
                    img = stuff['content']['img']
                    label = stuff['content']['label']
                    source = stuff['content']['source']
                    caption = stuff['content']['caption']

                    notebook += '<center id="c'+str(cell)+'"><form id="F'+str(cell)+'" enctype="multipart/form-data" method="POST" action="image" style="display: none;"><input type="file" name="img" value="Images/'+img+'" id="'+str(cell)+'" style="display: none;"><br>Label:<input name="label" value="'+label+'" id="L'+str(cell)+'"><br>Caption:<input name="caption" value="'+caption+'" id="C'+str(cell)+'"><br>Source:<input name="source" value="'+source+'" id="S'+str(cell)+'"></form></center>'
                    output = stuff['output']
                    notebook += '<center tabindex="0" id="co'+str(cell)+'"><div id="o'+str(cell)+'" class="paragraph">'+output+'</div></center>'
            else:
                content = stuff['content']
                output = stuff['output']
                notebook += '<center id="c'+str(cell)+'"><textarea id="'+str(cell)+'" action="evalCell" style="width: 800px; display: none;">'+content+'</textarea></center>'
                notebook += '<center tabindex="0" id="co'+str(cell)+'"><div id="o'+str(cell)+'" class="paragraph">'+output+'</div></center>'
        return notebook


if __name__ == '__main__':
    conf = {'/':{'tools.sessions.on':True,'tools.staticdir.on':True,'tools.staticdir.dir':os.getcwd()}}
    cherrypy.quickstart(WillNotebook(),'/',config=conf)
