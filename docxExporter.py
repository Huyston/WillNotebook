from docx import Document
from docx.shared import Inches,Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def getAllInside(first,last,content):
    # returns a dict of keys first+value+last and the values
    valuesDict = {}
    for i in range(content.count(first)):
        try:
            f = content.index(first)
            l = content[f+len(first):].index(last)
            key = content[f:f+len(first)+l+len(last)]
            value = key.replace(first,'').replace(last,'')
        except ValueError:
            print('Error')
            continue
        valuesDict[key] = value
        content = content.replace(key,'')
    return valuesDict

def getInside(first,last,content):
    f = content.index(first)
    l = content[f+len(first):].index(last)
    key = content[f:f+len(first)+l+len(last)]
    value = key.replace(first,'').replace(last,'')
    return value

class DocxExporter():
    def __init__(self,filename,docID,docType):
        self.docID = docID
        self.filename = filename
        if docType == 'abntepusp':
            self.document = Document(docx=os.getcwd()+'/Modelos/abntepusp.docx')
        else:
            self.document = Document()
        self.docType = docType
        self.docProperties = {}
        self.figNumber = 1
        self.tabNumber = 1
        self.inBullet = False

    def formatText(self,text,p):
        '''handle bold, italic, etc '''
        '''content = text
        if '<b>' in text:
            toBold = getAllInside('<b>','</b>',text)
            for bold in toBold:
                text = text.replace(bold,'\\textbf{'+toBold[bold]+'}')
            content = text
        if '<i>' in text:
            toItalicize = getAllInside('<i>','</i>',text)
            for italic in toItalicize:
                text = text.replace(italic,'\\textit{'+toItalicize[italic]+'}')
            content = text
        if '<code><pre>' in text: #multiline code
            codes = getAllInside('<code><pre>','</pre></code>',text)
            for code in codes:
                text = text.replace(code,'\\begin{lstlisting}\n'+codes[code]+'\n\end{lstlisting}')
            content = text
        if '<code>' in text: #inline code
            codes = getAllInside('<code>','</code>',text)
            for code in codes:
                text = text.replace(code,'\lstinline{'+codes[code]+'}')
            content = text
        return content'''
        if text:
            textBuffer = ''
            italic = False
            bold = False
            for letter in text:
                textBuffer += letter
                textBuffer = textBuffer.replace('\n',' ')
                if '<i>' in textBuffer:
                    runText = textBuffer.replace('<i>','')
                    r = p.add_run(runText)
                    r.italic = italic
                    r.bold = bold
                    textBuffer = ''
                    italic = True
                elif '</i>' in textBuffer:
                    runText = textBuffer.replace('</i>','')
                    r = p.add_run(runText)
                    r.italic = italic
                    r.bold = bold
                    textBuffer = ''
                    italic = False
                elif '<b>' in textBuffer:
                    runText = textBuffer.replace('<b>','')
                    r = p.add_run(runText)
                    r.italic = italic
                    r.bold = bold
                    textBuffer = ''
                    bold = True
                elif '</b>' in textBuffer:
                    runText = textBuffer.replace('</b>','')
                    r = p.add_run(runText)
                    r.italic = italic
                    r.bold = bold
                    textBuffer = ''
                    bold = False
                elif '</p>' in textBuffer:
                    print('entrou no ref')
                    ### This happens when a reference is written
                    ### Normally the reference is between </a> and </p>
                    reference = getInside('</a>','</p>',textBuffer).replace('&#x00A0;',' ').replace('&#8211;','-')
                    print(reference)
                    refBuffer = ''
                    p = self.document.add_paragraph()
                    for letter in reference:
                        refBuffer += letter
                        if 'class="ecti-1000">' in refBuffer:
                            initEmphasis = '<span'+getInside('<span','class="ecti-1000">',refBuffer)+'class="ecti-1000">'
                            runText = refBuffer.replace(initEmphasis,'')
                            r = p.add_run(runText)
                            r.italic = italic
                            r.bold = bold
                            refBuffer = ''
                            bold = True
                            print('Adding non bold')
                        elif '</span>' in refBuffer:
                            runText = refBuffer.replace('</span>','')
                            r = p.add_run(runText)
                            r.italic = italic
                            r.bold = bold
                            refBuffer = ''
                            bold = False
                            print('Adding bold')
                    if refBuffer:
                        p.add_run(refBuffer)
                    textBuffer = ''

            if textBuffer:
                p.add_run(textBuffer)

    def addText(self,text):
        p = self.document.add_paragraph()
        self.formatText(text,p)

    def addHeading(self,title,level,label):
        p = self.document.add_heading(level=level)
        self.formatText(title,p)

    def addFigure(self,img,caption,source='',label='',width='0.5'):

        cap = self.document.add_paragraph('Fig. '+str(self.figNumber)+': '+caption,'Caption')
        cap.alignment = 1
        para = self.document.add_paragraph()
        run = para.add_run()
        run.add_picture(os.getcwd()+'/Archieves/'+self.docID+'/Images/'+img,width=int(float(width)*Inches(5.75)))
        para.alignment = 1
        sou = self.document.add_paragraph('Source: '+source,'Caption')
        sou.alignment = 1
        self.figNumber += 1
        self.document.add_paragraph() #pula linha

    def addTable(self,table,caption,label='',source=''):
        cols = 0
        for row in table.split('\n'):
            rowCols = 0
            if '||' in row:
                rowCols = len(row.split('||'))
            elif '|' in row:
                rowCols = len(row.split('|'))
            if rowCols > cols:
                cols = rowCols

        cap = self.document.add_paragraph('Table '+str(self.tabNumber)+': '+caption,'Caption')
        cap.alignment = 1
        para = self.document.add_paragraph()
        docxTable = self.document.add_table(rows=1,cols=cols)
        docxTable.alignment = 1
        for row in table.split('\n'):
            if '||' in row:
                headings = row.split('||')
                hCells = docxTable.add_row().cells
                for n,heading in enumerate(headings):
                    hCells[n].text = heading
            elif '|' in row:
                colText = row.split('|')
                hCells = docxTable.add_row().cells
                for n,text in enumerate(colText):
                    hCells[n].text = text
        sou = self.document.add_paragraph('Source: '+source,'Caption')
        sou.alignment = 1
        self.tabNumber += 1
        self.document.add_paragraph() #pula linha

    def addTitle(self,title):
        if self.docType == 'abntepusp':
            self.docProperties['title'] = title

    def addAuthor(self,author):
        if self.docType == 'abntepusp':
            self.docProperties['author'] = author

    def addAdvisor(self,advisor):
        if self.docType == 'abntepusp':
            self.docProperties['advisor'] = advisor

    def addConcentrationArea(self,area):
        if self.docType == 'abntepusp':
            self.docProperties['concentration'] = area

    def addDepartment(self,department):
        if self.docType == 'abntepusp':
            self.docProperties['department'] = department

    def addModelType(self,modelType,area):
        if self.docType == 'abntepusp':
            self.docProperties['modelType'] = modelType
            self.docProperties['area'] = area

    def addLocal(self,local):
        if self.docType == 'abntepusp':
            self.docProperties['local'] = local

    def addDate(self,date):
        if self.docType == 'abntepusp':
            self.docProperties['date'] = date

    def makeCover(self):
        def set_row_height(row, height):
            height = float(height)
            height = height*10000.0/176.4
            tr = row._tr
            trPr = tr.find(qn('w:trPr'))
            if trPr==None:
                x=OxmlElement('w:trPr')
                tr.append(x);
                trPr = tr.find(qn('w:trPr'));
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(height))
            trPr.append(trHeight)

        def cover(coverType,titleText,author,advisor,concentration,department,modelType,area,local,dateText):
            #template
            tab = self.document.add_table(cols=2, rows=4, style='SimpleTable2')
            set_row_height(tab.rows[0],90)
            set_row_height(tab.rows[1],30)
            set_row_height(tab.rows[2],108)
            set_row_height(tab.rows[3],10)
            authorName = tab.cell(0,0).merge(tab.cell(0,1))
            authorName.width = Mm(210-50)
            authorName.text = author
            authorName.paragraphs[0].style = 'Normal'
            authorName.paragraphs[0].bold = True
            authorName.paragraphs[0].alignment = 1
            authorName.paragraphs[0].first_line_indent = None
            title = tab.cell(1,0).merge(tab.cell(1,1))
            title.width = Mm(210-50)
            title.text = titleText
            title.paragraphs[0].style = 'Body Text' # Precisa ser o Body text 3 ... 
            title.paragraphs[0].alignment = 1
            apresBlank = tab.cell(2,0)
            apresBlank.width = Mm((210-50)/2.0)
            apres = tab.cell(2,1)
            apres.width = Mm((210-50)/2.0)
            apres.text = 'Documento bla'
            apres.paragraphs[0].style = 'Normal'
            apres.paragraphs[0].bold = True
            apres.paragraphs[0].alignment = 3
            apres.paragraphs[0].first_line_indent = None
            date = tab.cell(3,0).merge(tab.cell(3,1))
            date.width = Mm(210-50)
            date.text = local+'\n'+dateText
            date.paragraphs[0].style = 'Normal'
            date.paragraphs[0].bold = True
            date.paragraphs[0].alignment = 1

        if self.docType == 'abntepusp':
            cover('capa',
                  self.docProperties['title'],
                  self.docProperties['author'],
                  self.docProperties['advisor'],
                  self.docProperties['concentration'],
                  self.docProperties['department'],
                  self.docProperties['modelType'],
                  self.docProperties['area'],
                  self.docProperties['local'],
                  self.docProperties['date'],
                  )
            cover('falsafolhaderosto',
                  self.docProperties['title'],
                  self.docProperties['author'],
                  self.docProperties['advisor'],
                  self.docProperties['concentration'],
                  self.docProperties['department'],
                  self.docProperties['modelType'],
                  self.docProperties['area'],
                  self.docProperties['local'],
                  self.docProperties['date'],
                  )
            cover('folhaderosto',
                  self.docProperties['title'],
                  self.docProperties['author'],
                  self.docProperties['advisor'],
                  self.docProperties['concentration'],
                  self.docProperties['department'],
                  self.docProperties['modelType'],
                  self.docProperties['area'],
                  self.docProperties['local'],
                  self.docProperties['date'],
                  )

    def addBullet(self,topic):
        self.document.add_paragraph(topic,style='List Bullet') 

    def close(self):
        self.document.save(os.getcwd()+'/Archieves/'+self.docID+'/'+self.filename+'.docx')


